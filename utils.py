import os
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st # For st.error in PDF export if fpdf is missing
# Conditional import for fpdf will be handled within the export_to_pdf function
import unicodedata



# Helper function to remove problematic Unicode characters
def remove_problematic_chars(text):
    """Removes characters that might cause encoding or display issues,
       especially those outside common encodings like latin-1, by replacing
       common problematic characters and filtering others."""
    if not isinstance(text, str):
        return text # Return as is if not a string

    # Replace common problematic characters with safe alternatives
    cleaned_text = text.replace('\u2013', '-') # En dash
    cleaned_text = cleaned_text.replace('\u2014', '-') # Em dash
    cleaned_text = cleaned_text.replace('\u2018', "'") # Left single quote
    cleaned_text = cleaned_text.replace('\u2019', "'") # Right single quote (apostrophe)
    cleaned_text = cleaned_text.replace('\u201c', '"') # Left double quote
    cleaned_text = cleaned_text.replace('\u201d', '"') # Right double quote
    cleaned_text = cleaned_text.replace('\u2026', '...') # Ellipsis
    cleaned_text = cleaned_text.replace('\u2022', '*') # Bullet point
    cleaned_text = cleaned_text.replace('\u2122', '(TM)') # Trade Mark symbol
    cleaned_text = cleaned_text.replace('\u00AE', '(R)') # Registered symbol
    cleaned_text = cleaned_text.replace('\u00A9', '(C)') # Copyright symbol


    # Attempt standard encode/decode with 'ignore' errors for characters still problematic
    # This is a fallback and might lose some characters not explicitly handled above
    try:
        # Try encoding to latin-1 and then decoding. Characters not in latin-1 will be ignored.
        cleaned_text = cleaned_text.encode('latin-1', errors='ignore').decode('latin-1')
    except UnicodeEncodeError:
        # If latin-1 encoding still fails unexpectedly, try a more aggressive UTF-8 encode/decode with replace
        cleaned_text = cleaned_text.encode('utf-8', errors='replace').decode('utf-8')
    except Exception as e:
        # Catch any other potential encoding errors during this step
        print(f"Warning: Error during initial encoding cleanup: {e}. Proceeding with filtering.")


    # Further filter potentially problematic Unicode characters using a stricter regex
    # This regex aims to keep printable ASCII, common whitespace, and a limited set of safe Unicode.
    # It removes control characters, surrogates, and many other non-ASCII characters.
    # Keeping only printable ASCII and basic whitespace:
    # \x20-\x7E : Printable ASCII characters (space to tilde)
    # \n\r\t : Newline, carriage return, tab
    # \u00A0 : Non-breaking space (often useful)
    # \u20AC : Euro sign (common) - can add others if needed
    # This regex is stricter than the previous one to avoid latin-1 issues.
    # Adjust the regex based on what characters are expected/allowed in the output documents.
    # Current regex keeps printable ASCII, newline, carriage return, tab, and non-breaking space.
    problematic_chars_regex = re.compile(r'[^\x20-\x7E\n\r\t\u00A0]')


    filtered_text = problematic_chars_regex.sub('', cleaned_text)

    return filtered_text


# Load configuration
def load_config():
    """Load configuration from config.json or create default if not exists"""
    config_path = "config.json"

    default_config = {
        "company_info": {
            "name": "Your Company Name",
            "logo_path": "",
            "default_styles": {
                "primary_color": "#003366",
                "secondary_color": "#669933",
                "font_family": "Arial"
            }
        },
        "api_keys": {
            "openai_key": os.environ.get("OPENAI_API_KEY", ""),
            "google_api_key": os.environ.get("GOOGLE_API_KEY", "")
        },
        "knowledge_base": {
            "directory": "markdown_responses",
            "embedding_model": "all-MiniLM-L6-v2",
            "metadata_fields": ["client_industry", "proposal_success", "project_size", "key_differentiators"]
        },
        "proposal_settings": {
            "default_sections": [],
            "max_tokens_per_section": 2000,
            "templates": ["Standard RFP", "Technical RFP", "Commercial RFP"]
        },
        "internal_capabilities": {
            "technical": ["Cloud solutions", "AI implementation", "Data analytics"],
            "functional": ["Project management", "24/7 support", "Custom development"]
        },
        "scoring_system": {
            "weighting": {
                "requirement_match": 0.4,
                "compliance": 0.25,
                "quality": 0.2,
                "alignment": 0.15,
                "risk": 0.1
            },
            "grading_scale": {
                "excellent": [90, 100],
                "good": [70, 89],
                "fair": [50, 69],
                "poor": [0, 49]
            }
        }
    }

    if not os.path.exists(config_path):
        print("config.json not found, creating default.")
        with open(config_path, 'w') as f:
            json.dump(default_config, f, indent=4)
        return default_config

    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
            # Merge with default config to ensure all keys exist
            # This handles cases where config.json exists but is missing sections
            merged_config = default_config.copy()
            merged_config.update(config)
            return merged_config
    except json.JSONDecodeError:
        print(f"Error decoding JSON from {config_path}. Using default config.")
        return default_config
    except Exception as e:
        print(f"An unexpected error occurred loading {config_path}: {e}. Using default config.")
        return default_config
    
    

# Word export function
def export_to_word(proposal_data, company_name, client_name, output_path, company_logo_path=None):
    """Export the generated proposal to a professionally formatted Word document"""
    doc = Document()

    # Set document styles
    styles = doc.styles

    # Modify heading styles
    heading1 = styles['Heading 1']
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = styles['Heading 2']
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    # Set document properties
    doc.core_properties.author = remove_problematic_chars(company_name) if company_name else ""
    doc.core_properties.title = remove_problematic_chars(f"Proposal for {client_name}") if client_name else "Proposal"


    # Add title page
    if company_logo_path and os.path.exists(company_logo_path):
        try:
            doc.add_picture(company_logo_path, width=Inches(2.0))
            doc.add_paragraph()  # Add some space after logo
        except Exception as e:
            print(f"Could not add logo to document: {e}")


    # Ensure client_name is cleaned before adding to document
    cleaned_client_name = remove_problematic_chars(client_name) if client_name else "Client"
    title = doc.add_heading(f"Proposal for {cleaned_client_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph()
    # Ensure company_name is cleaned before adding to document
    cleaned_company_name = remove_problematic_chars(company_name) if company_name else "Your Company Name"
    subtitle_run = subtitle.add_run(f"Prepared by {cleaned_company_name}")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a page break
    doc.add_page_break()

    # Add table of contents title
    doc.add_heading("Table of Contents", 1)

    # Generate table of contents
    # Note: Accurate TOC page numbers in docx are complex and usually require a second pass
    # or using built-in docx features which are harder to control programmatically.
    # This provides a basic list of sections.
    toc = doc.add_paragraph()
    for idx, section_name in enumerate(proposal_data["sections"]):
        # Ensure section_name is cleaned for TOC
        cleaned_section_name = remove_problematic_chars(section_name)
        toc.add_run(f"{cleaned_section_name}").bold = True
        toc.add_run(f"... [Page Number]\n") # Placeholder


    # Add page break after TOC
    doc.add_page_break()

    # Add each section with proper formatting
    for section_name, section_content in proposal_data["sections"].items():
        # Ensure section_name is cleaned for heading
        cleaned_section_name = remove_problematic_chars(section_name)
        doc.add_heading(cleaned_section_name, 1)

        # Ensure section_content is cleaned before processing lines
        cleaned_section_content = remove_problematic_chars(section_content)
        lines = cleaned_section_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if line.startswith('### '):
                # Ensure heading text is cleaned
                doc.add_heading(remove_problematic_chars(line[4:].strip()), 3)
            elif line.startswith('## '):
                # Ensure heading text is cleaned
                doc.add_heading(remove_problematic_chars(line[3:].strip()), 2)
            elif line.startswith('# '):
                # Ensure heading text is cleaned
                doc.add_heading(remove_problematic_chars(line[2:].strip()), 1)
            elif line.startswith('- ') or line.startswith('* '):
                # Ensure list item text is cleaned
                p = doc.add_paragraph(remove_problematic_chars(line[2:]), style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                # Ensure list item text is cleaned
                p = doc.add_paragraph(remove_problematic_chars(re.sub(r'^\d+\.\s', '', line)), style='List Number')
            elif line.startswith('|') and i+1 < len(lines) and '|--' in lines[i+1]:
                # Basic table parsing
                table_rows = []
                table_rows.append(line)
                i += 1
                while i < len(lines) and lines[i].startswith('|'):
                    table_rows.append(lines[i])
                    i += 1
                if len(table_rows) > 1: # Need at least header and one data row (or just header if parsing allows)
                    # Assuming header is the first row and separator is the second
                    header_cells = [remove_problematic_chars(cell.strip()) for cell in table_rows[0].split('|')[1:-1]] # Clean header cells
                    num_cols = len(header_cells)
                    if num_cols > 0:
                        # Count data rows (excluding header and separator)
                        data_rows = [row for row in table_rows[2:] if row.strip() and '|' in row]
                        num_rows = len(data_rows) + 1 # Add 1 for the header row

                        if num_rows > 0:
                             table = doc.add_table(rows=num_rows, cols=num_cols)
                             table.style = 'Table Grid'

                             # Add header row
                             for j, cell_text in enumerate(header_cells):
                                 table.cell(0, j).text = cell_text # Header cells are already cleaned

                             # Add data rows
                             for row_idx, row_text in enumerate(data_rows):
                                 cells = [remove_problematic_chars(cell.strip()) for cell in row_text.split('|')[1:-1]] # Clean data cells
                                 for j, cell_text in enumerate(cells):
                                     if j < num_cols: # Ensure we don't go out of bounds
                                         table.cell(row_idx+1, j).text = cell_text # Data cells are already cleaned
                        else:
                             # Handle case with only a header and separator, no data rows
                             table = doc.add_table(rows=1, cols=num_cols)
                             table.style = 'Table Grid'
                             for j, cell_text in enumerate(header_cells):
                                 table.cell(0, j).text = cell_text # Header cells are already cleaned

                i -= 1 # Decrement i because the while loop incremented it past the table
            elif line:
                # Ensure paragraph text is cleaned
                p = doc.add_paragraph(remove_problematic_chars(line))
            i += 1

        if section_name != list(proposal_data["sections"].keys())[-1]:
            doc.add_page_break()

    doc.save(output_path)

    return output_path

# PDF export function
def export_to_pdf(proposal_data, company_name, client_name, output_path, company_logo_path=None):
    try:
        from fpdf import FPDF
    except ImportError:
        st.error("The 'fpdf' library is required for PDF export. Please install it: pip install fpdf")
        return None

    class ProposalPDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 12)
            # Ensure client_name is cleaned before adding to header
            client_name_str = remove_problematic_chars(str(client_name)) if client_name else "Client"
            self.cell(0, 10, txt=f"Proposal for {client_name_str}", border=0, ln=1, align='C')

            if company_logo_path and os.path.exists(company_logo_path):
                try:
                    self.image(company_logo_path, 10, 8, 30)
                except Exception as e:
                    print(f"Could not add logo to PDF: {e}")

            self.ln(20)

        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, txt=f'Page {self.page_no()}', border=0, ln=0, align='C')

    pdf = ProposalPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", size=12)

    # Table of Contents
    pdf.cell(200, 10, txt="Table of Contents", ln=True, align='C')
    pdf.ln(5)

    # Note: Calculating accurate page numbers in PDF based on content length is complex.
    # This provides a basic estimate.
    current_page = 3 # Start after Title and TOC pages
    for section_name in proposal_data["sections"]:
         # Ensure section_name is cleaned for TOC
         cleaned_section_name = remove_problematic_chars(section_name)
         pdf.cell(0, 10, txt=f"{cleaned_section_name} - Page {current_page}", ln=True)
         # Estimate pages for the next section - a rough estimate
         # This is highly dependent on font size, line height, page margins, etc.
         # A more accurate method would involve rendering the content and counting pages.
         # Ensure content used for estimation is cleaned
         content = remove_problematic_chars(proposal_data["sections"][section_name])
         lines_per_page_estimate = 40 # Rough estimate
         estimated_lines = len(content.split('\n'))
         estimated_pages = max(1, estimated_lines // lines_per_page_estimate)
         current_page += estimated_pages


    pdf.add_page()

    for section_name, content in proposal_data["sections"].items():
        # Ensure section_name is cleaned for heading
        cleaned_section_name = remove_problematic_chars(section_name)
        pdf.set_font("Arial", 'B', 16)
        pdf.multi_cell(0, 10, txt=cleaned_section_name, border=0) # Use multi_cell for long titles
        pdf.ln(5) # Reduced line break after section title

        pdf.set_font("Arial", size=12)
        # Ensure content is cleaned before splitting into lines
        cleaned_content = remove_problematic_chars(content)
        lines = cleaned_content.split('\n')
        for line in lines:
            # Handle basic markdown like bold/italic if needed, fpdf requires specific commands
            # For simplicity here, just print lines. More complex formatting requires parsing markdown.
            # Ensure line is cleaned before printing
            cleaned_line = remove_problematic_chars(line)
            cleaned_line = re.sub(r'[\*_`]', '', cleaned_line) # Remove basic markdown chars
            if cleaned_line.strip(): # Avoid adding empty lines
                 pdf.multi_cell(0, 6, txt=cleaned_line, border=0)
                 pdf.ln(1) # Reduced line break between paragraphs

        # Add page break if it's not the last section
        if section_name != list(proposal_data["sections"].keys())[-1]:
            pdf.add_page()

    pdf.output(output_path)
    return output_path