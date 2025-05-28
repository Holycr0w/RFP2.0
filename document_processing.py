import re
from docx import Document
import PyPDF2
from utils import remove_problematic_chars


# Document processing functions
def extract_text_from_docx(file_path):
    """Extract text from DOCX files including tables and headers"""
    doc = Document(file_path)
    full_text = []

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Apply cleaning to cell text
                cleaned_cell_text = remove_problematic_chars(cell.text.strip())
                if cleaned_cell_text:
                    row_text.append(cleaned_cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))

    for para in doc.paragraphs:
        # Apply cleaning to paragraph text
        cleaned_para_text = remove_problematic_chars(para.text.strip())
        if cleaned_para_text:
            if para.style.name.startswith('Heading'):
                heading_level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                prefix = '#' * heading_level + ' '
                full_text.append(f"{prefix}{cleaned_para_text}")
            else:
                full_text.append(cleaned_para_text)

    return '\n'.join(full_text) # Text is already cleaned


def extract_text_from_pdf(file_path):
    """Extract text from PDF documents"""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = []
        for page in reader.pages:
            # Apply cleaning to extracted page text
            cleaned_page_text = remove_problematic_chars(page.extract_text())
            text.append(cleaned_page_text)
    return '\n'.join(text) # Text is already cleaned


def extract_sections_from_rfp(rfp_text):
    """Extract structured sections from the RFP text with improved pattern matching"""
    # Ensure the input text is cleaned before processing
    cleaned_rfp_text = remove_problematic_chars(rfp_text)

    section_patterns = [
        r'^(?:\d+\.)?(?:\d+\.)?(?:\d+\.)?\s*([A-Z][A-Za-z\s]+)$',
        r'^([A-Z][A-Z\s]+)(?:\:|\.)?\s*$',
        r'^(?:Section|SECTION)\s+\d+\s*[\:\-\.]\s*([A-Za-z\s]+)$'
    ]

    sections = {}
    current_section = "Overview"
    current_content = []

    for line in cleaned_rfp_text.split('\n'):
        matched = False
        for pattern in section_patterns:
            match = re.match(pattern, line.strip())
            if match:
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                    current_content = []

                current_section = match.group(1).strip()
                matched = True
                break

        if not matched:
            current_content.append(line)

    if current_content:
        sections[current_section] = '\n'.join(current_content)

    return sections

def process_rfp(file_path):
    """Extract text from uploaded RFP document"""
    if file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.md') or file_path.endswith('.txt'):
        # Added errors='replace' to handle problematic characters during reading
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
        return remove_problematic_chars(content) # Clean content after reading
    else:
        raise ValueError("Unsupported file format. Please use DOCX, PDF, TXT or MD file.")